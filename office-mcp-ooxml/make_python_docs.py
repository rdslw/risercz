from pathlib import Path
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.chart import BarChart, Reference

OUT = Path(__file__).parent / 'artifacts'
OUT.mkdir(exist_ok=True)

def make_docx():
    doc = Document()
    doc.add_heading('Acme Cloud Platform P&L FY2026', 0)
    doc.add_paragraph('Scenario: Board-ready operating plan with revenue, COGS, opex, EBITDA, risks, and sales/PM actions.')
    doc.add_heading('Executive summary', 1)
    for text in [
        'Revenue grows from $12.0M to $17.7M on enterprise expansion and partner attach.',
        'Gross margin improves from 68.0% to 73.0% after support automation and hosting optimization.',
        'EBITDA turns from -$1.4M in Q1 to $2.0M in Q4 while preserving PM discovery capacity.',
    ]:
        doc.add_paragraph(text, style='List Bullet')
    doc.add_heading('Quarterly P&L', 1)
    headers = ['Metric', 'Q1', 'Q2', 'Q3', 'Q4', 'FY']
    rows = [
        ['Revenue', 12000000, 13800000, 15500000, 17700000, 59000000],
        ['COGS', -3840000, -3864000, -4185000, -4779000, -16668000],
        ['Gross Profit', 8160000, 9936000, 11315000, 12921000, 42332000],
        ['Sales & Marketing', -4800000, -5100000, -5350000, -5700000, -20950000],
        ['Product & Engineering', -3500000, -3600000, -3700000, -3850000, -14650000],
        ['G&A', -1300000, -1350000, -1400000, -1420000, -5470000],
        ['EBITDA', -1440000, -114000, 865000, 1951000, 1262000],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row): cells[i].text = str(val)
    doc.add_heading('Commentary', 1)
    doc.add_paragraph('Sales should prioritize expansion SKUs with payback under 12 months; PM should protect roadmap items that lift attach and reduce support load.')
    doc.add_heading('Risk register', 1)
    r = doc.add_table(rows=1, cols=3); r.style = 'Medium Grid 1 Accent 2'
    for i,h in enumerate(['Risk','Impact','Mitigation']): r.rows[0].cells[i].text = h
    for row in [['Enterprise slips','-$1.8M revenue','Stage-gate late-quarter commits'],['Cloud inflation','-220 bps GM','Reserved capacity + workload rightsizing'],['Feature churn','Delayed attach','PM scorecard tied to conversion']]:
        c=r.add_row().cells
        for i,v in enumerate(row): c[i].text=v
    doc.save(OUT / 'python_pl.docx')

def make_xlsx():
    wb = Workbook(); ws = wb.active; ws.title = 'P&L'
    data = [
        ['Metric','Q1','Q2','Q3','Q4','FY'],
        ['Revenue',12000000,13800000,15500000,17700000,'=SUM(B2:E2)'],
        ['COGS',-3840000,-3864000,-4185000,-4779000,'=SUM(B3:E3)'],
        ['Gross Profit','=B2+B3','=C2+C3','=D2+D3','=E2+E3','=SUM(B4:E4)'],
        ['Gross Margin','=B4/B2','=C4/C2','=D4/D2','=E4/E2','=F4/F2'],
        ['Sales & Marketing',-4800000,-5100000,-5350000,-5700000,'=SUM(B6:E6)'],
        ['Product & Engineering',-3500000,-3600000,-3700000,-3850000,'=SUM(B7:E7)'],
        ['G&A',-1300000,-1350000,-1400000,-1420000,'=SUM(B8:E8)'],
        ['EBITDA','=B4+SUM(B6:B8)','=C4+SUM(C6:C8)','=D4+SUM(D6:D8)','=E4+SUM(E6:E8)','=SUM(B9:E9)'],
    ]
    for row in data: ws.append(row)
    for cell in ws[1]: cell.font = Font(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='1F4E78')
    for row in ws.iter_rows(min_row=2, min_col=2):
        for cell in row: cell.number_format = '0.0%' if cell.row == 5 else '$#,##0;[Red]($#,##0)'
    ws.freeze_panes='B2'; ws.auto_filter.ref='A1:F9'
    ws['A11']='PM/Sales action'; ws['B11']='Prioritize expansion play + support automation before Q3.'
    chart=BarChart(); chart.title='Revenue vs EBITDA'; chart.add_data(Reference(ws,min_col=2,max_col=5,min_row=2,max_row=2), titles_from_data=False); chart.add_data(Reference(ws,min_col=2,max_col=5,min_row=9,max_row=9), titles_from_data=False); chart.set_categories(Reference(ws,min_col=2,max_col=5,min_row=1)); ws.add_chart(chart,'H2')
    wb.save(OUT / 'python_pl.xlsx')

make_docx(); make_xlsx(); print(OUT)
